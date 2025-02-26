from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
import io

app = Flask(__name__)
CORS(app, resources={r"/generate-doc": {"origins": "http://localhost:3000"}})

@app.route("/generate-doc", methods=["POST"])
def generate_doc():
    try:
        data = request.json
        name = data.get("name", "Unknown")
        date = data.get("date", "Unknown")
        comment = data.get("comment", "No comment")

        doc = Document()
        doc.add_heading("Generated Document", level=1)
        doc.add_paragraph(f"Name: {name}")
        doc.add_paragraph(f"Date: {date}")
        doc.add_paragraph(f"Comment: {comment}")

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return send_file(
            doc_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="document.docx"
        )
    except Exception as e:
        return {"error": str(e)}, 500

if __name__ == '__main__':
    app.run(debug=True)