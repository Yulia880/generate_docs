from docx import Document
import os

def create_word_document(data):
    doc = Document()
    doc.add_heading("Generated Document", level=1)

    doc.add_paragraph(f"Name: {data.get('name', '')}")
    doc.add_paragraph(f"Date: {data.get('date', '')}")
    doc.add_paragraph(f"Comment:\n{data.get('name', '')}")

    file_path = "output.docx"
    doc.save(file_path)
    return file_path